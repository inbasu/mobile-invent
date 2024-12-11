from docx import Document


class WordDocument:
    def __init__(self, path: str) -> None:
        self.document = Document(path)

    

 

    def change_table(self, data: dict[str, str]):
        if not data:
            return self.document
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        if not data:
                            return self.document
                        if par.text in data.keys():
                            par.text = data.pop(par.text)
        return self.document

    
    def save(self, path) -> None:
        self.document.save(path)
